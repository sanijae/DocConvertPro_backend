# from io import BytesIO
# import PyPDF2
# import zlib

# def ccompress_pdf(input_pdf_path, output_pdf_path):
#     with open(input_pdf_path, 'rb') as input_file:
#         pdf_reader = PyPDF2.PdfReader(input_file)
#         pdf_writer = PyPDF2.PdfWriter()

#         # Copy pages from input PDF to output PDF
#         for page in pdf_reader.pages:
#             pdf_writer.add_page(page)

#         # Write the PDF content to a BytesIO object
#         output_buffer = BytesIO()
#         pdf_writer.write(output_buffer)
#         pdf_content = output_buffer.getvalue()

#         # Compress the content using zlib
#         #compressed_content = zlib.compress(pdf_content)

#         # Write the compressed content to the output file
#         with open(output_pdf_path, 'wb') as output_file:
#             output_file.write(pdf_content)

# import PyPDF2

# def compress_pdfs(input_pdf_path, output_pdf_path):
#     with open(input_pdf_path, 'rb') as input_file:
#         pdf_reader = PyPDF2.PdfReader(input_file)

#         # Create a PDF writer object
#         pdf_writer = PyPDF2.PdfWriter()

#         # Copy pages from input PDF to output PDF
#         for page in pdf_reader.pages:
#             page.compress_content_streams()  # This is CPU intensive!
#             pdf_writer.add_page(page)

#         # Write the compressed PDF to the output file
#         with open(output_pdf_path, 'wb') as output_file:
#             pdf_writer.write(output_file)

# import subprocess

# def compress_pdf_with_ghostscript(input_pdf_path, output_pdf_path):
#     # Command to compress PDF using ghostscript
#     command = [
#         'gs',  # Ghostscript command
#         '-sDEVICE=pdfwrite',  # Output device
#         '-dCompatibilityLevel=1.4',
#         '-dNOPAUSE',
#         '-dQUIET',
#         '-dBATCH',
#         '-dPDFSETTINGS=/printer',  # Compression level (change as needed)
#         '-o', output_pdf_path,  # Output file
#         input_pdf_path  # Input file
#     ]

#     # Execute the ghostscript command
#     subprocess.run(command)

# # Example usage
# input_pdf_path = "file.pdf"
# output_pdf_path = "d_file.pdf"

# compress_pdf_with_ghostscript(input_pdf_path, output_pdf_path)


# # Example usage
# # input_pdf = "reduced_pdf.pdf"
# # output_pdf = "reduceded_pdf.pdf"
# #ccompress_pdf(input_pdf, output_pdf)


from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

def cconvert_docx_to_pdf(docx_file, pdf_file):
    # Create a PDF document
    pdf = SimpleDocTemplate(pdf_file, pagesize=letter)

    # Styles for paragraphs
    styles = getSampleStyleSheet()
    
    # Additional styles for document title and headers
    title_style = ParagraphStyle(name='Title', parent=styles['Title'])
    header_style = ParagraphStyle(name='Header', parent=styles['Heading1'])

    # Container for the elements
    content = []

    # Load the DOCX file
    doc = Document(docx_file)

    # Iterate through each element in the DOCX file
    for element in doc.element.body:
        if element.tag.endswith('p'):  # Paragraph
            text = element.text
            # Check if element has a style and if it's a title or header
            if element.style and element.style.name == 'Title':
                content.append(Paragraph(text, title_style))
            elif element.style and element.style.name.startswith('Heading'):
                content.append(Paragraph(text, header_style))
            else:
                content.append(Paragraph(text, styles['Normal']))
            content.append(Spacer(1, 12))  # Add some space between paragraphs
        elif element.tag.endswith('tbl'):  # Table
            table_data = []
            for row in element.iterchildren():
                table_row = [cell.text for cell in row.iterchildren()]
                table_data.append(table_row)
            content.append(Table(table_data))
            content.append(Spacer(1, 12))  # Add some space after table
        elif element.tag.endswith('drawing'):  # Drawing (Image)
            image_part = element.find('.//pic/blipFill/blip')
            if image_part is not None:
                image_filename = image_part.attrib.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if image_filename:
                    image_path = 'word/media/' + image_filename
                    content.append(Image(image_path))
                    content.append(Spacer(1, 12))  # Add some space after image
        elif element.tag.endswith('object'):  # Embedded Object (Figure)
            # Handle embedded objects here (if necessary)
            pass

    # Build PDF document
    pdf.build(content)

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table
from reportlab.lib.styles import getSampleStyleSheet

def convert_docx_to_pdf(docx_file, pdf_file):
    # Create a PDF document
    pdf = SimpleDocTemplate(pdf_file, pagesize=letter)

    # Styles for paragraphs
    styles = getSampleStyleSheet()

    # Container for the elements
    content = []

    # Load the DOCX file
    doc = Document(docx_file)

    # Extract and add paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text
        content.append(Paragraph(text, styles['Normal']))
        content.append(Spacer(1, 12))  # Add some space between paragraphs

    # Extract and add tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_row = [cell.text for cell in row.cells]
            table_data.append(table_row)
        content.append(Table(table_data))
        content.append(Spacer(1, 12))  # Add some space after table

    # Extract and add images
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            image_data = rel.target_part.blob
            content.append(Image(image_data))
            content.append(Spacer(1, 12))  # Add some space after image

    # Build PDF document
    pdf.build(content)


# Usage
inputFile = "word.docx"
outputFile = "wmw_To_pdf.pdf"
convert_docx_to_pdf(inputFile, outputFile)
